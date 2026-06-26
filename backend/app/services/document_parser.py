from abc import ABC, abstractmethod
from pathlib import Path
import fitz  # pymupdf

from app.core.exceptions import DocumentProcessingError, UnsupportedFileTypeError


class DocumentParser(ABC):
    """Abstract base class for document parsers (Strategy Pattern)."""

    @abstractmethod
    def parse(self, file_path: Path) -> tuple[str, int]:
        pass

    @abstractmethod
    def supports(self, extension: str) -> bool:
        pass


class PDFParser(DocumentParser):
    """Parses PDF files using PyMuPDF — text-based PDFs only."""

    def parse(self, file_path: Path) -> tuple[str, int]:
        try:
            doc = fitz.open(str(file_path))
            pages = []
            page_count = len(doc)

            for page in doc:
                text = page.get_text("text").strip()
                if text:
                    pages.append(text)

            doc.close()

            if not pages:
                raise DocumentProcessingError(
                    "This PDF appears to be scanned (image-only) and contains no "
                    "selectable text. Please upload a text-based PDF or a DOCX file."
                )

            return "\n\n".join(pages), page_count

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse PDF: {e}") from e

    def supports(self, extension: str) -> bool:
        return extension.lower() == ".pdf"


class DOCXParser(DocumentParser):
    """Parses DOCX files using python-docx."""

    def parse(self, file_path: Path) -> tuple[str, int]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            full_text = "\n\n".join(paragraphs)
            if not full_text.strip():
                raise DocumentProcessingError("DOCX appears to be empty.")

            word_count = len(full_text.split())
            estimated_pages = max(1, word_count // 500)
            return full_text, estimated_pages

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse DOCX: {e}") from e

    def supports(self, extension: str) -> bool:
        return extension.lower() == ".docx"


class DocumentParserFactory:
    """Factory for selecting the appropriate parser (Factory Pattern)."""

    def __init__(self):
        self._parsers: list[DocumentParser] = [
            PDFParser(),
            DOCXParser(),
        ]

    def get_parser(self, file_path: Path) -> DocumentParser:
        extension = file_path.suffix.lower()
        for parser in self._parsers:
            if parser.supports(extension):
                return parser
        raise UnsupportedFileTypeError(
            f"File type '{extension}' is not supported. Please upload PDF or DOCX."
        )